"""Page layout, header/footer, spacing, bookmark, and watermark tools.

Cross-platform tools using python-docx for section-level and page-level operations.
"""

import json
import os
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


async def set_page_layout(
    filename: str,
    section_index: int = 0,
    orientation: str = None,
    page_width_inches: float = None,
    page_height_inches: float = None,
    margin_top_inches: float = None,
    margin_bottom_inches: float = None,
    margin_left_inches: float = None,
    margin_right_inches: float = None,
) -> str:
    """Set page layout for a document section.

    Args:
        filename: Path to the Word document.
        section_index: Section number (0-based). Default 0 = first section.
        orientation: "portrait" or "landscape".
        page_width_inches: Page width in inches.
        page_height_inches: Page height in inches.
        margin_top_inches: Top margin in inches.
        margin_bottom_inches: Bottom margin in inches.
        margin_left_inches: Left margin in inches.
        margin_right_inches: Right margin in inches.
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, err = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {err}"

    try:
        doc = Document(filename)
        if section_index >= len(doc.sections):
            return f"Section {section_index} does not exist. Document has {len(doc.sections)} section(s)."

        section = doc.sections[section_index]
        changes = []

        if orientation is not None:
            if orientation.lower() == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                # Swap width/height if needed
                if section.page_width < section.page_height:
                    section.page_width, section.page_height = section.page_height, section.page_width
                changes.append("orientation=landscape")
            elif orientation.lower() == "portrait":
                section.orientation = WD_ORIENT.PORTRAIT
                if section.page_width > section.page_height:
                    section.page_width, section.page_height = section.page_height, section.page_width
                changes.append("orientation=portrait")

        if page_width_inches is not None:
            section.page_width = Inches(page_width_inches)
            changes.append(f"width={page_width_inches}in")
        if page_height_inches is not None:
            section.page_height = Inches(page_height_inches)
            changes.append(f"height={page_height_inches}in")
        if margin_top_inches is not None:
            section.top_margin = Inches(margin_top_inches)
            changes.append(f"margin_top={margin_top_inches}in")
        if margin_bottom_inches is not None:
            section.bottom_margin = Inches(margin_bottom_inches)
            changes.append(f"margin_bottom={margin_bottom_inches}in")
        if margin_left_inches is not None:
            section.left_margin = Inches(margin_left_inches)
            changes.append(f"margin_left={margin_left_inches}in")
        if margin_right_inches is not None:
            section.right_margin = Inches(margin_right_inches)
            changes.append(f"margin_right={margin_right_inches}in")

        doc.save(filename)
        return json.dumps({
            "success": True,
            "section": section_index,
            "changes": changes,
        })

    except Exception as e:
        return json.dumps({"error": str(e)})


async def add_header_footer(
    filename: str,
    section_index: int = 0,
    header_text: str = None,
    footer_text: str = None,
    header_alignment: str = "center",
    footer_alignment: str = "center",
) -> str:
    """Add header and/or footer text to a document section.

    Args:
        filename: Path to the Word document.
        section_index: Section number (0-based).
        header_text: Text to put in the header. None = don't change.
        footer_text: Text to put in the footer. None = don't change.
        header_alignment: "left", "center", "right", "justify".
        footer_alignment: "left", "center", "right", "justify".
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, err = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {err}"

    try:
        doc = Document(filename)
        if section_index >= len(doc.sections):
            return f"Section {section_index} does not exist."

        section = doc.sections[section_index]
        added = []

        if header_text is not None:
            header = section.header
            header.is_linked_to_previous = False
            # Clear existing and add new
            for p in header.paragraphs:
                p.clear()
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.text = header_text
            p.alignment = _ALIGN_MAP.get(header_alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)
            added.append("header")

        if footer_text is not None:
            footer = section.footer
            footer.is_linked_to_previous = False
            for p in footer.paragraphs:
                p.clear()
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.text = footer_text
            p.alignment = _ALIGN_MAP.get(footer_alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)
            added.append("footer")

        doc.save(filename)
        return json.dumps({"success": True, "added": added, "section": section_index})

    except Exception as e:
        return json.dumps({"error": str(e)})


def _add_field(paragraph, field_code: str):
    """Insert a Word field code (PAGE, NUMPAGES, etc.) into a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    run._r.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> {field_code} </w:instrText>'
    )
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    )
    run3._r.append(fld_char_end)


async def add_page_numbers(
    filename: str,
    section_index: int = 0,
    position: str = "footer",
    alignment: str = "center",
    prefix: str = "",
    suffix: str = "",
    include_total: bool = False,
) -> str:
    """Add page numbers to a document.

    Args:
        filename: Path to the Word document.
        section_index: Section number (0-based).
        position: "header" or "footer".
        alignment: "left", "center", "right".
        prefix: Text before page number (e.g. "Page ").
        suffix: Text after page number.
        include_total: If True, adds " of N" after the page number.
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, err = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {err}"

    try:
        doc = Document(filename)
        if section_index >= len(doc.sections):
            return f"Section {section_index} does not exist."

        section = doc.sections[section_index]
        target = section.header if position == "header" else section.footer
        target.is_linked_to_previous = False

        # Add a new paragraph for page numbers
        p = target.add_paragraph()
        p.alignment = _ALIGN_MAP.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)

        if prefix:
            p.add_run(prefix)

        _add_field(p, "PAGE")

        if include_total:
            p.add_run(" / ")
            _add_field(p, "NUMPAGES")

        if suffix:
            p.add_run(suffix)

        doc.save(filename)
        return json.dumps({
            "success": True,
            "position": position,
            "alignment": alignment,
            "include_total": include_total,
        })

    except Exception as e:
        return json.dumps({"error": str(e)})


async def add_section_break(
    filename: str,
    break_type: str = "new_page",
) -> str:
    """Add a section break to the end of the document.

    Args:
        filename: Path to the Word document.
        break_type: "new_page", "continuous", "even_page", "odd_page".
    """
    from docx.enum.section import WD_SECTION

    type_map = {
        "new_page": WD_SECTION.NEW_PAGE,
        "continuous": WD_SECTION.CONTINUOUS,
        "even_page": WD_SECTION.EVEN_PAGE,
        "odd_page": WD_SECTION.ODD_PAGE,
    }

    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, err = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {err}"

    if break_type not in type_map:
        return f"Invalid break_type: {break_type}. Use: {list(type_map.keys())}"

    try:
        doc = Document(filename)
        doc.add_section(type_map[break_type])
        doc.save(filename)
        return json.dumps({
            "success": True,
            "break_type": break_type,
            "total_sections": len(doc.sections),
        })

    except Exception as e:
        return json.dumps({"error": str(e)})


async def set_paragraph_spacing(
    filename: str,
    paragraph_index: int = None,
    start_paragraph: int = None,
    end_paragraph: int = None,
    space_before_pt: float = None,
    space_after_pt: float = None,
    line_spacing: float = None,
    line_spacing_rule: str = None,
) -> str:
    """Set paragraph spacing for one or a range of paragraphs.

    Args:
        filename: Path to the Word document.
        paragraph_index: Single paragraph (0-based). Ignored if start/end given.
        start_paragraph: Start of range (inclusive, 0-based).
        end_paragraph: End of range (inclusive, 0-based).
        space_before_pt: Space before paragraph in points.
        space_after_pt: Space after paragraph in points.
        line_spacing: Line spacing value (depends on rule).
        line_spacing_rule: "single", "1.5_lines", "double", "exactly", "at_least", "multiple".
    """
    from docx.enum.text import WD_LINE_SPACING

    rule_map = {
        "single": WD_LINE_SPACING.SINGLE,
        "1.5_lines": WD_LINE_SPACING.ONE_POINT_FIVE,
        "double": WD_LINE_SPACING.DOUBLE,
        "exactly": WD_LINE_SPACING.EXACTLY,
        "at_least": WD_LINE_SPACING.AT_LEAST,
        "multiple": WD_LINE_SPACING.MULTIPLE,
    }

    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, err = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {err}"

    try:
        doc = Document(filename)
        total = len(doc.paragraphs)

        # Determine range
        if start_paragraph is not None and end_paragraph is not None:
            indices = range(start_paragraph, min(end_paragraph + 1, total))
        elif paragraph_index is not None:
            if paragraph_index >= total:
                return f"Paragraph {paragraph_index} does not exist. Document has {total} paragraphs."
            indices = [paragraph_index]
        else:
            # Apply to all
            indices = range(total)

        count = 0
        for i in indices:
            pf = doc.paragraphs[i].paragraph_format
            if space_before_pt is not None:
                pf.space_before = Pt(space_before_pt)
            if space_after_pt is not None:
                pf.space_after = Pt(space_after_pt)
            if line_spacing_rule is not None and line_spacing_rule in rule_map:
                pf.line_spacing_rule = rule_map[line_spacing_rule]
            if line_spacing is not None:
                if line_spacing_rule in ("exactly", "at_least"):
                    pf.line_spacing = Pt(line_spacing)
                else:
                    pf.line_spacing = line_spacing
            count += 1

        doc.save(filename)
        return json.dumps({
            "success": True,
            "paragraphs_affected": count,
        })

    except Exception as e:
        return json.dumps({"error": str(e)})


async def add_bookmark(
    filename: str,
    paragraph_index: int,
    bookmark_name: str,
) -> str:
    """Add a named bookmark at a paragraph.

    Args:
        filename: Path to the Word document.
        paragraph_index: Paragraph to bookmark (0-based).
        bookmark_name: Bookmark name (no spaces, alphanumeric + underscore).
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, err = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {err}"

    try:
        doc = Document(filename)
        if paragraph_index >= len(doc.paragraphs):
            return f"Paragraph {paragraph_index} does not exist."

        para = doc.paragraphs[paragraph_index]

        # Generate a unique bookmark ID
        import random
        bm_id = str(random.randint(1000, 99999))

        # Insert bookmarkStart before paragraph content
        bm_start = parse_xml(
            f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="{bookmark_name}"/>'
        )
        bm_end = parse_xml(
            f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>'
        )

        para._p.insert(0, bm_start)
        para._p.append(bm_end)

        doc.save(filename)
        return json.dumps({
            "success": True,
            "bookmark_name": bookmark_name,
            "paragraph_index": paragraph_index,
        })

    except Exception as e:
        return json.dumps({"error": str(e)})


async def add_watermark(
    filename: str,
    text: str = "TASLAK",
    font_size: int = 72,
    font_color: str = "C0C0C0",
    rotation: int = -45,
    section_index: int = 0,
) -> str:
    """Add a diagonal text watermark to a document.

    Args:
        filename: Path to the Word document.
        text: Watermark text (e.g. "DRAFT", "TASLAK", "GİZLİ").
        font_size: Font size in points.
        font_color: Hex color without # (e.g. "C0C0C0").
        rotation: Rotation angle in degrees (e.g. -45).
        section_index: Which section to add watermark to (0-based).
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, err = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {err}"

    try:
        doc = Document(filename)
        if section_index >= len(doc.sections):
            return f"Section {section_index} does not exist."

        section = doc.sections[section_index]
        header = section.header
        header.is_linked_to_previous = False

        # Build a VML shape for the watermark
        # The shape is placed in a paragraph inside the header
        watermark_xml = (
            f'<w:r {nsdecls("w", "v", "o", "w10")}>'
            f'  <w:rPr><w:noProof/></w:rPr>'
            f'  <w:pict>'
            f'    <v:shapetype id="_x0000_t136" coordsize="21600,21600"'
            f'      o:spt="136" adj="10800"'
            f'      path="m@7,l@8,m@5,21600l@6,21600e">'
            f'      <v:formulas>'
            f'        <v:f eqn="sum #0 0 10800"/>'
            f'        <v:f eqn="prod #0 2 1"/>'
            f'        <v:f eqn="sum 21600 0 @1"/>'
            f'        <v:f eqn="sum 0 0 @2"/>'
            f'        <v:f eqn="sum 21600 0 @3"/>'
            f'        <v:f eqn="if @0 @3 0"/>'
            f'        <v:f eqn="if @0 21600 @1"/>'
            f'        <v:f eqn="if @0 0 @2"/>'
            f'        <v:f eqn="if @0 @4 21600"/>'
            f'        <v:f eqn="mid @5 @6"/>'
            f'        <v:f eqn="mid @8 @5"/>'
            f'        <v:f eqn="mid @7 @8"/>'
            f'        <v:f eqn="mid @6 @7"/>'
            f'        <v:f eqn="sum @6 0 @5"/>'
            f'      </v:formulas>'
            f'      <v:path textpathok="t" o:connecttype="custom"'
            f'        o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800"'
            f'        o:connectangles="270,180,90,0"/>'
            f'      <v:textpath on="t" fitshape="t"/>'
            f'      <v:handles><v:h position="#0,bottomRight" xrange="6629,14971"/></v:handles>'
            f'      <o:lock v:ext="edit" text="t" shapetype="t"/>'
            f'    </v:shapetype>'
            f'    <v:shape id="PowerPlusWaterMarkObject"'
            f'      o:spid="_x0000_s2049" type="#_x0000_t136"'
            f'      style="position:absolute;margin-left:0;margin-top:0;'
            f'width:500pt;height:150pt;rotation:{rotation};'
            f'z-index:-251658752;mso-position-horizontal:center;'
            f'mso-position-horizontal-relative:margin;'
            f'mso-position-vertical:center;'
            f'mso-position-vertical-relative:margin"'
            f'      o:allowincell="f"'
            f'      fillcolor="#{font_color}" stroked="f">'
            f'      <v:fill opacity=".5"/>'
            f'      <v:textpath style="font-family:&quot;Calibri&quot;;font-size:{font_size}pt"'
            f'        string="{text}"/>'
            f'      <w10:wrap anchorx="margin" anchory="margin"/>'
            f'    </v:shape>'
            f'  </w:pict>'
            f'</w:r>'
        )

        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p._p.append(parse_xml(watermark_xml))

        doc.save(filename)
        return json.dumps({
            "success": True,
            "text": text,
            "font_size": font_size,
            "color": font_color,
            "rotation": rotation,
        })

    except Exception as e:
        return json.dumps({"error": str(e)})
